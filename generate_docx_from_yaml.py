#!/usr/bin/env python3
"""
Word Document Generator from YAML Resume Data
Generates both formatted and ATS-friendly versions from structured content
"""

import yaml
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class WordResumeGenerator:
    """Generates Word documents from YAML resume data"""

    def __init__(self, content: dict, style: str = 'formatted'):
        """
        Initialize generator

        Args:
            content: Resume content dictionary from YAML
            style: 'formatted' for styled resume, 'ats' for machine-readable
        """
        self.content = content
        self.style = style
        self.doc = Document()
        self.setup_document()

    def setup_document(self):
        """Configure document styles and margins"""
        # Set margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Configure styles
        styles = self.doc.styles

        # Normal text
        style_normal = styles['Normal']
        style_normal.font.name = 'Calibri'
        style_normal.font.size = Pt(11)

        # Heading 1 (Name)
        style_h1 = styles['Heading 1']
        style_h1.font.name = 'Calibri'
        style_h1.font.size = Pt(20 if self.style == 'formatted' else 16)
        style_h1.font.bold = True
        style_h1.font.color.rgb = RGBColor(0, 0, 0)

        # Heading 2 (Section headers)
        style_h2 = styles['Heading 2']
        style_h2.font.name = 'Calibri'
        style_h2.font.size = Pt(14)
        style_h2.font.bold = True
        style_h2.font.color.rgb = RGBColor(0, 0, 0)

    def add_horizontal_line(self):
        """Add a horizontal line for section separation (formatted style only)"""
        if self.style == 'formatted':
            p = self.doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.space_before = Pt(2)
            p_format.space_after = Pt(2)

            # Add bottom border
            pPr = p._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '7A1D0A')
            pBdr.append(bottom)
            pPr.append(pBdr)

    def add_header(self, project_dir=None):
        """Add resume header with name and contact info"""
        contact = self.content.get('contact', {})
        photo = contact.get('photo', '')

        # If photo exists, create a table for layout (photo + text)
        if photo and self.style == 'formatted' and project_dir:
            # Try to find the photo file
            photo_path = None
            for ext in ['.jpg', '.jpeg', '.png', '.JPG', '.JPEG', '.PNG']:
                potential_path = project_dir / (photo + ext)
                if potential_path.exists():
                    photo_path = potential_path
                    break

            if photo_path and photo_path.exists():
                # Create table for layout: photo on right, text on left
                table = self.doc.add_table(rows=1, cols=2)
                table.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Left cell - contact info
                left_cell = table.rows[0].cells[0]
                left_cell.width = Inches(4.5)

                # Name in left cell
                name_para = left_cell.paragraphs[0]
                name_run = name_para.add_run(contact.get('name', 'Your Name'))
                name_run.font.size = Pt(20)
                name_run.font.bold = True

                # Contact info in left cell
                contact_para = left_cell.add_paragraph()
                contact_items = []
                if contact.get('email'):
                    contact_items.append(contact['email'])
                if contact.get('phone'):
                    contact_items.append(contact['phone'])
                if contact.get('location'):
                    contact_items.append(contact['location'])
                if contact.get('linkedin'):
                    contact_items.append(f"LinkedIn: {contact['linkedin']}")
                if contact.get('github'):
                    contact_items.append(f"GitHub: {contact['github']}")

                contact_run = contact_para.add_run(' | '.join(contact_items))
                contact_run.font.size = Pt(10)

                # Right cell - photo
                right_cell = table.rows[0].cells[1]
                right_cell.width = Inches(1.5)
                photo_para = right_cell.paragraphs[0]
                photo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                # Add photo with size constraint
                photo_run = photo_para.add_run()
                photo_run.add_picture(str(photo_path), width=Inches(1.2))

                # Hide table borders for cleaner look
                for row in table.rows:
                    for cell in row.cells:
                        tcPr = cell._element.get_or_add_tcPr()
                        tcBorders = OxmlElement('w:tcBorders')
                        for border in ['top', 'left', 'bottom', 'right']:
                            border_el = OxmlElement(f'w:{border}')
                            border_el.set(qn('w:val'), 'none')
                            tcBorders.append(border_el)
                        tcPr.append(tcBorders)
            else:
                # Photo specified but not found, fall back to text-only
                self._add_text_only_header(contact)
        else:
            # No photo or ATS mode - text only
            self._add_text_only_header(contact)

        self.add_horizontal_line()

    def _add_text_only_header(self, contact):
        """Add text-only header (no photo)"""
        # Name
        heading = self.doc.add_heading(contact.get('name', 'Your Name'), level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contact info
        contact_para = self.doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        contact_items = []
        if contact.get('email'):
            contact_items.append(contact['email'])
        if contact.get('phone'):
            contact_items.append(contact['phone'])
        if contact.get('location'):
            contact_items.append(contact['location'])
        if contact.get('linkedin'):
            contact_items.append(f"LinkedIn: {contact['linkedin']}")
        if contact.get('github'):
            contact_items.append(f"GitHub: {contact['github']}")

        contact_run = contact_para.add_run(' | '.join(contact_items))
        if self.style == 'formatted':
            contact_run.font.size = Pt(10)

    def add_summary(self):
        """Add professional summary section"""
        summary = self.content.get('summary', '')
        if not summary:
            return

        self.doc.add_heading('Professional Summary', level=2)
        self.doc.add_paragraph(summary.strip())
        self.add_horizontal_line()

    def add_employment(self):
        """Add employment section"""
        jobs = self.content.get('employment', [])
        if not jobs:
            return

        self.doc.add_heading('Professional Experience', level=2)

        for job in jobs:
            # Job header
            job_para = self.doc.add_paragraph()

            # Title
            title_run = job_para.add_run(job.get('title', ''))
            title_run.bold = True
            title_run.font.size = Pt(12)

            # Company and dates
            job_para.add_run(f" | {job.get('company', '')}")

            date_para = self.doc.add_paragraph()
            date_run = date_para.add_run(job.get('dates', ''))
            date_run.italic = True
            if self.style == 'formatted':
                date_run.font.size = Pt(10)

            # Summary (if exists)
            summary = job.get('summary', '')
            if summary:
                summary_para = self.doc.add_paragraph(summary)
                summary_para.paragraph_format.left_indent = Inches(0.25)
                if self.style == 'formatted':
                    summary_para.runs[0].italic = True
                    summary_para.runs[0].font.size = Pt(10)

            # Achievements
            for achievement in job.get('achievements', []):
                if self.style == 'ats':
                    # Plain bullets for ATS
                    bullet = self.doc.add_paragraph(achievement, style='List Bullet')
                    bullet.paragraph_format.left_indent = Inches(0.25)
                else:
                    # Formatted bullets
                    bullet = self.doc.add_paragraph(achievement, style='List Bullet')
                    bullet.paragraph_format.left_indent = Inches(0.25)
                    bullet.paragraph_format.space_after = Pt(2)

            # Add spacing between jobs
            self.doc.add_paragraph()

        self.add_horizontal_line()

    def add_education(self):
        """Add education section"""
        degrees = self.content.get('education', [])
        if not degrees:
            return

        self.doc.add_heading('Education', level=2)

        for degree in degrees:
            edu_para = self.doc.add_paragraph()

            # Degree name
            degree_run = edu_para.add_run(degree.get('degree', ''))
            degree_run.bold = True
            degree_run.font.size = Pt(11)

            # Details
            details_para = self.doc.add_paragraph()
            details_text = f"{degree.get('dates', '')} | {degree.get('details', '')}"
            details_run = details_para.add_run(details_text)
            details_run.italic = True
            if self.style == 'formatted':
                details_run.font.size = Pt(10)

            # Achievements (if any)
            for achievement in degree.get('achievements', []):
                bullet = self.doc.add_paragraph(achievement, style='List Bullet')
                bullet.paragraph_format.left_indent = Inches(0.25)

        self.add_horizontal_line()

    def add_projects(self):
        """Add technical projects section"""
        projects = self.content.get('projects', [])
        if not projects:
            return

        self.doc.add_heading('Technical Projects', level=2)

        for project in projects:
            proj_para = self.doc.add_paragraph()

            # Project name
            name_run = proj_para.add_run(project.get('name', ''))
            name_run.bold = True

            # Description
            description = project.get('description', '')
            if description:
                proj_para.add_run(f": {description}")

            # URL (if exists)
            url = project.get('url', '')
            if url:
                url_para = self.doc.add_paragraph()
                url_run = url_para.add_run(f"URL: {url}")
                url_run.italic = True
                if self.style == 'formatted':
                    url_run.font.size = Pt(9)
                    url_run.font.color.rgb = RGBColor(122, 29, 10)  # MarkerColour

        self.add_horizontal_line()

    def add_skills(self):
        """Add technical skills section"""
        skills = self.content.get('skills', {})
        if not skills:
            return

        self.doc.add_heading('Technical Skills', level=2)

        if self.style == 'ats':
            # ATS-friendly: Simple list format
            for category, skill_list in skills.items():
                skill_para = self.doc.add_paragraph()
                category_run = skill_para.add_run(f"{category}: ")
                category_run.bold = True

                if isinstance(skill_list, list):
                    skill_para.add_run(', '.join(skill_list))
                else:
                    skill_para.add_run(str(skill_list))
        else:
            # Formatted: Table layout for better visual organization
            table = self.doc.add_table(rows=len(skills), cols=2)
            table.style = 'Light Grid Accent 1'

            for idx, (category, skill_list) in enumerate(skills.items()):
                row = table.rows[idx]

                # Category cell
                category_cell = row.cells[0]
                category_para = category_cell.paragraphs[0]
                category_run = category_para.add_run(category)
                category_run.bold = True
                category_cell.width = Inches(1.5)

                # Skills cell
                skills_cell = row.cells[1]
                skills_para = skills_cell.paragraphs[0]

                if isinstance(skill_list, list):
                    skills_para.add_run(', '.join(skill_list))
                else:
                    skills_para.add_run(str(skill_list))

    def generate(self, output_path: Path, project_dir: Path = None):
        """Generate the complete Word document

        Args:
            output_path: Where to save the document
            project_dir: Directory containing the project files (for photo lookup)
        """
        self.add_header(project_dir=project_dir)
        self.add_summary()
        self.add_employment()
        self.add_education()
        self.add_projects()
        self.add_skills()

        # Save document
        self.doc.save(output_path)
        print(f"✓ Generated {self.style} Word document: {output_path}")


def main():
    """Main function to generate Word documents from YAML"""
    import sys

    # Parse command line arguments
    project_name = sys.argv[1] if len(sys.argv) > 1 else 'shelter-sre'
    mode = sys.argv[2] if len(sys.argv) > 2 else 'both'  # 'both', 'normal', or 'ats'

    # Paths
    PROJECT_DIR = Path(__file__).parent / 'projects' / project_name
    CONTENT_FILE = PROJECT_DIR / 'content.yaml'

    # Validate paths exist
    if not PROJECT_DIR.exists():
        print(f"Error: Project directory not found: {PROJECT_DIR}")
        sys.exit(1)

    if not CONTENT_FILE.exists():
        print(f"Error: content.yaml not found: {CONTENT_FILE}")
        sys.exit(1)

    # Load content
    with open(CONTENT_FILE, 'r') as f:
        content = yaml.safe_load(f)

    if mode == 'both' or mode == 'normal':
        # Generate formatted version
        formatted_output = PROJECT_DIR / 'resume.docx'
        formatted_gen = WordResumeGenerator(content, style='formatted')
        formatted_gen.generate(formatted_output, project_dir=PROJECT_DIR)

    if mode == 'both' or mode == 'ats':
        # Generate ATS-friendly version
        ats_output = PROJECT_DIR / 'resume-ats.docx'
        ats_gen = WordResumeGenerator(content, style='ats')
        ats_gen.generate(ats_output, project_dir=PROJECT_DIR)

    if mode == 'both':
        print(f"\n✓ Generated both versions successfully!")
        print(f"  - Formatted (styled): {PROJECT_DIR / 'resume.docx'}")
        print(f"  - ATS-friendly (plain): {PROJECT_DIR / 'resume-ats.docx'}")


if __name__ == '__main__':
    main()
