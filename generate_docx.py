#!/usr/bin/env python3
"""
Word Document Generator for LaTeX Resume
Converts LaTeX resume content to a machine-readable .docx format
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class ResumeDocxGenerator:
    """Generates a Word document from LaTeX resume files"""

    def __init__(self, project_dir: Path):
        self.project_dir = project_dir
        self.doc = Document()
        self.setup_styles()

    def setup_styles(self):
        """Configure document styles"""
        # Set narrow margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

    def clean_latex(self, text: str) -> str:
        """Remove LaTeX commands and convert to plain text"""
        # Remove common LaTeX commands
        text = re.sub(r'\\tech\{([^}]+)\}', r'\1', text)
        text = re.sub(r'\\metric\{([^}]+)\}', r'\1', text)
        text = re.sub(r'\\company\{([^}]+)\}', r'\1', text)
        text = re.sub(r'\\role\{([^}]+)\}', r'\1', text)
        text = re.sub(r'\\textbf\{([^}]+)\}', r'\1', text)
        text = re.sub(r'\\emph\{([^}]+)\}', r'\1', text)
        text = re.sub(r'\\small', '', text)
        text = re.sub(r'\\Large', '', text)
        text = re.sub(r'\\relsize\{[^}]+\}', '', text)
        text = re.sub(r'\\vspace\{[^}]+\}', '', text)
        text = re.sub(r'\\achievement\{([^}]+)\}', r'\1', text)
        text = re.sub(r'\\LaTeX', 'LaTeX', text)
        text = re.sub(r'\\&', '&', text)
        text = re.sub(r'\\_', '_', text)
        text = re.sub(r'\\pcidss\{\}', 'PCI-DSS', text)
        text = re.sub(r'\\vocalinkbold\{\}', 'Vocalink', text)
        text = re.sub(r'\\mdesbold\{\}', 'MDES', text)
        text = re.sub(r'\\billions', 'billions', text)

        # Remove remaining LaTeX commands
        text = re.sub(r'\\[a-zA-Z]+\{([^}]+)\}', r'\1', text)
        text = re.sub(r'\\[a-zA-Z]+', '', text)

        # Clean up spacing
        text = re.sub(r'\s+', ' ', text)
        text = text.strip()

        return text

    def add_header(self, name: str, contact_info: dict):
        """Add resume header with name and contact info"""
        # Name
        heading = self.doc.add_heading(name, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contact info
        contact_para = self.doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        contact_items = []
        if contact_info.get('email'):
            contact_items.append(f"Email: {contact_info['email']}")
        if contact_info.get('phone'):
            contact_items.append(f"Phone: {contact_info['phone']}")
        if contact_info.get('location'):
            contact_items.append(f"Location: {contact_info['location']}")
        if contact_info.get('linkedin'):
            contact_items.append(f"LinkedIn: {contact_info['linkedin']}")
        if contact_info.get('github'):
            contact_items.append(f"GitHub: {contact_info['github']}")

        contact_para.add_run(' | '.join(contact_items))

    def add_summary(self, summary_text: str):
        """Add professional summary section"""
        self.doc.add_heading('Professional Summary', level=2)
        summary = self.clean_latex(summary_text)
        self.doc.add_paragraph(summary)

    def parse_employment(self) -> list:
        """Parse employment.tex file"""
        employment_file = self.project_dir / 'employment.tex'
        content = employment_file.read_text()

        jobs = []
        # Match each job entry
        pattern = r'\\entry\*\[\\role\{([^}]+)\}\]([^|]+)\|\s*\\company\{([^}]+)\}(.*?)(?=\\entry\*|\\end\{rubric\})'

        for match in re.finditer(pattern, content, re.DOTALL):
            title = match.group(1)
            dates = match.group(2).strip()
            company = match.group(3)
            achievements_text = match.group(4)

            # Extract achievements
            achievements = []
            for achievement in re.findall(r'\\achievement\{([^}]+(?:\{[^}]+\}[^}]*)*)\}', achievements_text):
                cleaned = self.clean_latex(achievement)
                if cleaned:
                    achievements.append(cleaned)

            jobs.append({
                'title': title,
                'dates': dates,
                'company': company,
                'achievements': achievements
            })

        return jobs

    def parse_education(self) -> list:
        """Parse education.tex file"""
        education_file = self.project_dir / 'education.tex'
        content = education_file.read_text()

        degrees = []
        # Match each education entry
        pattern = r'\\entry\*\[\\textbf\{([^}]+)\}\]([^|]+)\|([^\n]+)'

        for match in re.finditer(pattern, content):
            degree = match.group(1)
            dates = match.group(2).strip()
            details = self.clean_latex(match.group(3))

            degrees.append({
                'degree': degree,
                'dates': dates,
                'details': details
            })

        return degrees

    def parse_skills(self) -> dict:
        """Parse skills.tex file"""
        skills_file = self.project_dir / 'skills.tex'
        content = skills_file.read_text()

        skills = {}
        # Match each skill category
        pattern = r'\\entry\*\[\\textbf\{([^}]+)\}\]\s*([^\n]+(?:\n(?!\\entry\*|\\vspace|\\end)[^\n]+)*)'

        for match in re.finditer(pattern, content):
            category = match.group(1)
            skills_text = self.clean_latex(match.group(2))
            skills[category] = skills_text

        return skills

    def parse_projects(self) -> list:
        """Parse projects.tex file"""
        projects_file = self.project_dir / 'projects.tex'
        content = projects_file.read_text()

        projects = []

        # Match githubproject entries
        pattern = r'\\githubproject\{([^}]+)\}\{([^}]+)\}\{([^}]+)\}'

        for match in re.finditer(pattern, content):
            name = self.clean_latex(match.group(1))
            description = self.clean_latex(match.group(2))
            url = match.group(3)

            projects.append({
                'name': name,
                'description': description,
                'url': url
            })

        return projects

    def add_employment_section(self, jobs: list):
        """Add employment section to document"""
        self.doc.add_heading('Professional Experience', level=2)

        for job in jobs:
            # Job title and company
            job_para = self.doc.add_paragraph()
            job_para.add_run(job['title']).bold = True
            job_para.add_run(f" | {job['company']}")
            job_para.add_run(f"\n{job['dates']}").italic = True

            # Achievements
            for achievement in job['achievements']:
                bullet = self.doc.add_paragraph(achievement, style='List Bullet')
                bullet.paragraph_format.left_indent = Inches(0.25)

            # Add spacing
            self.doc.add_paragraph()

    def add_education_section(self, degrees: list):
        """Add education section to document"""
        self.doc.add_heading('Education', level=2)

        for degree in degrees:
            edu_para = self.doc.add_paragraph()
            edu_para.add_run(degree['degree']).bold = True
            edu_para.add_run(f"\n{degree['dates']} | {degree['details']}").italic = True

    def add_projects_section(self, projects: list):
        """Add projects section to document"""
        self.doc.add_heading('Technical Projects', level=2)

        for project in projects:
            proj_para = self.doc.add_paragraph()
            proj_para.add_run(project['name']).bold = True
            proj_para.add_run(f": {project['description']}")
            if project['url']:
                proj_para.add_run(f"\nURL: {project['url']}").italic = True

    def add_skills_section(self, skills: dict):
        """Add skills section to document"""
        self.doc.add_heading('Technical Skills', level=2)

        for category, skills_text in skills.items():
            skill_para = self.doc.add_paragraph()
            skill_para.add_run(f"{category}: ").bold = True
            skill_para.add_run(skills_text)

    def generate(self, output_path: Path, contact_info: dict, summary: str):
        """Generate the complete Word document"""
        # Add header
        self.add_header(contact_info.get('name', 'Your Name'), contact_info)

        # Add summary
        if summary:
            self.add_summary(summary)

        # Parse and add sections
        jobs = self.parse_employment()
        self.add_employment_section(jobs)

        degrees = self.parse_education()
        self.add_education_section(degrees)

        projects = self.parse_projects()
        self.add_projects_section(projects)

        skills = self.parse_skills()
        self.add_skills_section(skills)

        # Save document
        self.doc.save(output_path)
        print(f"✓ Generated Word document: {output_path}")


def main():
    """Main function to generate Word document from LaTeX resume"""

    # Configuration
    PROJECT_DIR = Path(__file__).parent / 'projects' / 'shelter-sre'
    OUTPUT_PATH = Path(__file__).parent / 'projects' / 'shelter-sre' / 'resume.docx'

    # Contact information (update with your actual info)
    CONTACT_INFO = {
        'name': 'Adam E. Holbrook',
        'email': 'your.email@example.com',
        'phone': '(555) 123-4567',
        'location': 'St. Louis, MO',
        'linkedin': 'linkedin.com/in/yourprofile',
        'github': 'github.com/aeholbrook'
    }

    # Professional summary
    SUMMARY = """Senior Site Reliability Engineer with 5+ years in SRE/DevOps building reliable,
    observable infrastructure for mission-critical payment systems. Expert in CI/CD automation,
    Infrastructure-as-Code, and observability platforms (Splunk, Dynatrace, Domo).
    Proven ability to multiply team effectiveness through scalable tooling and operational excellence."""

    # Generate document
    generator = ResumeDocxGenerator(PROJECT_DIR)
    generator.generate(OUTPUT_PATH, CONTACT_INFO, SUMMARY)


if __name__ == '__main__':
    main()
